# -*- coding: utf-8 -*-
# @Time : 2023/3/5 11:28
# @Author : Boyang Li
# @File : genDocx.py
import os
import time
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt, Cm
from docx2pdf import convert

# 获取生成证明的日期（今日）
year = time.localtime().tm_year
month = time.localtime().tm_mon
day = time.localtime().tm_mday
month_parse = {1: 'January', 2: 'February', 3: 'March', 4: 'April', 5: 'May', 6: 'June',
               7: 'July', 8: 'August', 9: 'September', 10: 'October', 11: 'November', 12: 'December'}


def gen_document(text_zh: str, text_en: str, text_en2: str, stu_name: str, stu_num, path_, is_graduated, doc_type=1):
    """

    :param text_zh: 中文证明文本内容
    :param text_en: 英文证明文本内容
    :param text_en2: 英文第二段证明文本内容
    :param stu_name: 学生姓名（中文）
    :param stu_num: 学号
    :param path_: 输出文档路径
    :param is_graduated: 学生是否毕业
    :param doc_type: 输出文档类型。1：仅Word文档；2：仅PDF文档；3：两者均输出
    :return: None
    """

    document = Document()
    document.styles['Normal'].font.name = u'华文楷体'  # 设置文档的基础字体
    document.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'), u'华文楷体')  # 设置文档的基础中文字体
    document.styles['Normal'].font.size = Pt(14)  # 设置文档字体为14磅

    default_section = document.sections[0]
    # 默认宽度和高度（A4）
    default_section.page_width = Cm(21)
    default_section.page_height = Cm(29.7)
    p0 = document.add_paragraph()  # 加一行空行
    p0.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run0 = p0.add_run('')
    run0.font.size = Pt(22)
    run0.font.name = u'华文楷体'  # 设置西文字体
    run0.element.rPr.rFonts.set(qn('w:eastAsia'), u'华文楷体')  # 设置段中文字体

    # 建立中文标题行
    p1 = document.add_paragraph()  # 初始化
    p1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 对齐方式为居中
    run1 = p1.add_run('证          明')
    run1.font.name = u'华文楷体'  # 设置西文字体
    run1.element.rPr.rFonts.set(qn('w:eastAsia'), u'华文楷体')  # 设置段中文字体
    run1.font.size = Pt(24)  # 设置字体大小为24磅
    run1.font.bold = True  # 设置加粗
    p1.paragraph_format.space_before = Pt(10)  # 段前距离10磅

    # 第2段
    p2 = document.add_paragraph()  # 初始化
    run2 = p2.add_run(text_zh)  # 内容
    p2.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    run2.font.name = u'华文楷体'  # 设置西文字体
    run2.element.rPr.rFonts.set(qn('w:eastAsia'), u'华文楷体')  # 设置段中文字体
    run2.font.size = Pt(16)  # 设置字体大小为16磅
    p2.paragraph_format.first_line_indent = 406400
    p2.paragraph_format.line_spacing = Pt(36)

    # 第3段
    p3 = document.add_paragraph()
    run3 = p3.add_run('特此证明')  # 内容
    run3.font.name = u'华文楷体'  # 设置西文字体
    run3.element.rPr.rFonts.set(qn('w:eastAsia'), u'华文楷体')  # 设置段中文字体
    run3.font.size = Pt(16)  # 设置字体大小为16磅
    p3.paragraph_format.first_line_indent = 406400

    # 第4段
    p4 = document.add_paragraph()
    p4.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT  # 对齐方式为右对齐
    run4 = p4.add_run('上海师范大学商学院')  # 内容
    run4.font.name = u'华文楷体'  # 设置西文字体
    run4.element.rPr.rFonts.set(qn('w:eastAsia'), u'华文楷体')  # 设置段中文字体
    run4.font.size = Pt(16)  # 设置字体大小为16磅

    # 第5段
    p5 = document.add_paragraph()
    p5.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run5 = p5.add_run('%s年%s月%s日' % (year, month, day))  # 内容
    run5.font.name = u'华文楷体'  # 设置西文字体
    run5.element.rPr.rFonts.set(qn('w:eastAsia'), u'华文楷体')  # 设置段中文字体
    run5.font.size = Pt(16)  # 设置字体大小为16磅
    # run5.font.bold = True  # 加粗

    # 英文部分 标题段
    p6 = document.add_paragraph()  # 初始化段落
    p6.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 对齐方式为居中
    run6 = p6.add_run('Certificate')
    run6.font.name = u'华文楷体'  # 设置字体
    run6.font.size = Pt(24)  # 设置字体大小为24磅
    run6.font.bold = True  # 加粗

    # 称呼
    p7_0 = document.add_paragraph()
    run7_0 = p7_0.add_run('To whom it may concern, ')  # 内容
    p7_0.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY  # 双侧对齐
    run7_0.font.name = u'Times New Roman'  # 设置西文字体
    run7_0.font.size = Pt(14)  # 设置字体大小为14磅
    p7_0.paragraph_format.first_line_indent = 406400
    p7_0.paragraph_format.line_spacing = Pt(21)

    # 第7段
    p7 = document.add_paragraph()  # 初始化段落
    run7 = p7.add_run(text_en)  # 内容
    p7.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    run7.font.name = u'Times New Roman'  # 设置字体
    run7.font.size = Pt(14)  # 设置字体大小为14磅
    p7.paragraph_format.first_line_indent = 406400
    p7.paragraph_format.line_spacing = Pt(21)

    # 第8段
    p8 = document.add_paragraph()
    run8 = p8.add_run(text_en2)  # 内容
    p8.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    run8.font.name = u'Times New Roman'
    run8.font.size = Pt(14)  # 设置字体大小为16磅
    p8.paragraph_format.first_line_indent = 406400
    p8.paragraph_format.line_spacing = Pt(21)

    # 第9段
    p9 = document.add_paragraph()
    run9 = p9.add_run('Hereby certify.')  # 内容
    run9.font.name = u'Times New Roman'
    run9.font.size = Pt(14)  # 设置字体大小为16磅
    p9.paragraph_format.first_line_indent = 406400

    # 第10段
    p10 = document.add_paragraph()
    p10.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT  # 对齐方式为右侧对齐
    run10 = p10.add_run('School of Finance and Business')  # 内容
    run10.font.name = u'Times New Roman'
    run10.font.size = Pt(14)  # 设置字体大小为14磅

    # 第11段
    p11 = document.add_paragraph()
    p11.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT  # 对齐方式为右侧对齐
    run11 = p11.add_run('Shanghai Normal University')  # 内容
    run11.font.name = u'Times New Roman'
    run11.font.size = Pt(14)  # 设置字体大小为14磅

    # 第12段
    p12 = document.add_paragraph()
    p12.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT  # 对齐方式为右侧对齐
    run12 = p12.add_run('%s %s %s' % (day, month_parse[month], year))  # 内容
    run12.font.name = u'Times New Roman'
    run12.font.size = Pt(14)  # 设置字体大小为14磅

    # 保存到指定路径
    on_graduate = '_已毕业' if is_graduated else ''
    doc_path = path_ + '/%s_%s%s.docx' % (stu_num, stu_name, on_graduate)
    print(doc_type)
    document.save(doc_path)
    if doc_type == 2:  # 只保留PDF文档
        convert(doc_path, path_ + '/%s_%s%s.pdf' % (stu_num, stu_name, on_graduate))
        os.remove(doc_path)
    if doc_type == 3:  # 两种文档都保留
        convert(doc_path, path_ + '/%s_%s%s.pdf' % (stu_num, stu_name, on_graduate))

